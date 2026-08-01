from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .database import GENERATED_DIR, safe_filename, utc_now

APA_FONT = "Times New Roman"
APA_SIZE = Pt(12)


def _set_run_font(run, name: str = APA_FONT, size: Pt = APA_SIZE, bold: bool | None = None) -> None:
    run.font.name = name
    run.font.size = size
    if bold is not None:
        run.bold = bold
    # Word may ignore font.name for some scripts without these XML attributes.
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _set_apa_paragraph(paragraph, *, first_line: bool = True, align=None, space_after: Pt = Pt(0)) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    fmt.space_before = Pt(0)
    fmt.space_after = space_after
    if first_line:
        fmt.first_line_indent = Inches(0.5)
    if align is not None:
        paragraph.alignment = align


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    _set_run_font(run)


def _add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Haz clic derecho y elige Actualizar campo para generar el índice."
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(placeholder)
    run._r.append(fld_char3)
    _set_run_font(run)


def configure_apa_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = APA_FONT
    normal.font.size = APA_SIZE
    normal._element.rPr.rFonts.set(qn("w:ascii"), APA_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), APA_FONT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)

    for level in range(1, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = APA_FONT
        style.font.size = APA_SIZE
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)  # headings remain black, never Word blue
        style._element.rPr.rFonts.set(qn("w:ascii"), APA_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), APA_FONT)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.keep_with_next = True

    # APA-like levels: 1 centered; 2 and 3 left-aligned.
    doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.styles["Heading 3"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer.paragraphs[0]
    _add_page_number(footer)


def add_title_page(doc: Document, metadata: dict) -> None:
    for _ in range(4):
        p = doc.add_paragraph()
        _set_apa_paragraph(p, first_line=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_apa_paragraph(p, first_line=False)
    run = p.add_run(metadata.get("title") or "Monografía")
    _set_run_font(run, bold=True)

    fields = [
        metadata.get("author", ""),
        metadata.get("institution", "Universidad Tecnológica de los Andes"),
        metadata.get("course", ""),
        metadata.get("teacher", ""),
        metadata.get("city", "Abancay, Perú"),
        metadata.get("date", ""),
    ]
    for value in fields:
        if not value:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_apa_paragraph(p, first_line=False)
        run = p.add_run(str(value))
        _set_run_font(run)
    doc.add_page_break()


def _heading_level(line: str) -> tuple[int, str] | None:
    stripped = line.strip()
    if not stripped:
        return None
    markdown = re.match(r"^(#{1,3})\s+(.+)$", stripped)
    if markdown:
        return len(markdown.group(1)), markdown.group(2).strip()
    numbered = re.match(r"^(\d+(?:\.\d+){0,2})[.)]?\s+(.+)$", stripped)
    if numbered:
        level = min(numbered.group(1).count(".") + 1, 3)
        return level, f"{numbered.group(1)} {numbered.group(2).strip()}"
    upper = stripped.upper()
    known = {
        "INTRODUCCIÓN", "CONCLUSIONES", "RECOMENDACIONES", "REFERENCIAS",
        "BIBLIOGRAFÍA", "MARCO TEÓRICO", "METODOLOGÍA", "RESULTADOS",
        "DISCUSIÓN", "ANEXOS", "RESUMEN", "ABSTRACT",
    }
    if upper in known:
        return 1, stripped.title() if upper not in {"ABSTRACT"} else "Abstract"
    return None


def add_structured_text(doc: Document, text: str) -> None:
    blocks = re.split(r"\n\s*\n", (text or "").strip())
    for block in blocks:
        lines = [line.rstrip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        # A block containing only one heading is a heading.
        heading = _heading_level(lines[0]) if len(lines) == 1 else None
        if heading:
            level, title = heading
            p = doc.add_paragraph(style=f"Heading {level}")
            alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
            _set_apa_paragraph(p, first_line=False, align=alignment)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(title)
            _set_run_font(run, bold=True)
            continue

        # Bulleted and numbered lists.
        if all(re.match(r"^[-*•]\s+", line) for line in lines):
            for line in lines:
                p = doc.add_paragraph(style="List Bullet")
                _set_apa_paragraph(p, first_line=False)
                run = p.add_run(re.sub(r"^[-*•]\s+", "", line))
                _set_run_font(run)
            continue

        paragraph_text = " ".join(line.strip() for line in lines)
        p = doc.add_paragraph()
        _set_apa_paragraph(p, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        run = p.add_run(paragraph_text)
        _set_run_font(run)


def add_references(doc: Document, bibliography: str) -> None:
    if not bibliography.strip():
        return
    p = doc.add_paragraph(style="Heading 1")
    _set_apa_paragraph(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.keep_with_next = True
    run = p.add_run("Referencias")
    _set_run_font(run, bold=True)
    references = [item.strip() for item in re.split(r"\n\s*\n|\n(?=[A-ZÁÉÍÓÚÑ])", bibliography.strip()) if item.strip()]
    for reference in references:
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.left_indent = Inches(0.5)
        fmt.first_line_indent = Inches(-0.5)
        fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        fmt.space_after = Pt(0)
        run = p.add_run(" ".join(reference.splitlines()))
        _set_run_font(run)


def create_monograph_docx(metadata: dict, structured_text: str, bibliography: str) -> Path:
    """Crea una monografía APA 7 limpia, sin instrucciones internas visibles."""
    doc = Document()
    configure_apa_styles(doc)
    add_title_page(doc, metadata)
    add_structured_text(doc, structured_text)
    if bibliography.strip():
        references_section = doc.add_section(WD_SECTION.NEW_PAGE)
        references_section.top_margin = Inches(1)
        references_section.bottom_margin = Inches(1)
        references_section.left_margin = Inches(1)
        references_section.right_margin = Inches(1)
    add_references(doc, bibliography)
    name = safe_filename(metadata.get("title") or "monografia")
    target = GENERATED_DIR / safe_filename(f"{name}-{utc_now().replace(':', '-')}.docx")
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    return target


LEGAL_LABELS = {
    "solicitud": "SOLICITUD",
    "carta_notarial": "CARTA NOTARIAL",
    "apersonamiento": "ESCRITO DE APERSONAMIENTO",
    "reconsideracion": "RECURSO DE RECONSIDERACIÓN",
    "apelacion": "RECURSO DE APELACIÓN",
    "denuncia": "DENUNCIA",
    "poder_simple": "CARTA PODER SIMPLE",
    "descargo": "ESCRITO DE DESCARGO",
    "acta_reunion": "ACTA DE REUNIÓN",
}


def _configure_legal_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")


def _add_legal_paragraph(doc: Document, text: str, *, bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph = doc.add_paragraph()
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = 1.5
    if bold:
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    _set_run_font(run, name="Arial", size=Pt(11), bold=bold)
    return paragraph


def _create_minutes_docx(fields: dict, body: str) -> Path:
    """Crea un acta de reunión peruana con acuerdos y firmas editables."""
    doc = Document()
    _configure_legal_document(doc)

    title = _add_legal_paragraph(doc, "ACTA DE REUNIÓN", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.space_after = Pt(12)

    entity = str(fields.get("authority") or "").strip()
    responsible = str(fields.get("applicant") or "").strip()
    place = str(fields.get("address") or "").strip()
    city_date = str(fields.get("city_date") or "").strip()
    time_value = str(fields.get("time") or "").strip()
    purpose = str(fields.get("sumilla") or "").strip()
    participants = str(fields.get("participants") or "").strip()

    intro_parts = [f"En {place}", city_date]
    if time_value:
        intro_parts.append(f"a las {time_value}")
    intro = ", ".join(part for part in intro_parts if part)
    if entity:
        intro += f", se reunieron las personas vinculadas a {entity}"
    if responsible:
        intro += f", bajo la conducción de {responsible}"
    intro += "."
    _add_legal_paragraph(doc, intro)

    if purpose:
        _add_legal_paragraph(doc, "OBJETO DE LA REUNIÓN", bold=True)
        _add_legal_paragraph(doc, purpose)
    if participants:
        _add_legal_paragraph(doc, "PARTICIPANTES", bold=True)
        for participant in [item.strip() for item in participants.splitlines() if item.strip()]:
            _add_legal_paragraph(doc, participant, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    for heading, content in _split_legal_sections(body):
        _add_legal_paragraph(doc, heading.upper(), bold=True)
        for paragraph_text in content:
            paragraph = _add_legal_paragraph(doc, paragraph_text)
            paragraph.paragraph_format.first_line_indent = Inches(0.5)

    closing = "No habiendo otro asunto que tratar, se dio por concluida la reunión. Leída la presente acta, las personas participantes manifiestan su conformidad y la suscriben."
    _add_legal_paragraph(doc, closing)
    doc.add_page_break()
    _add_legal_paragraph(doc, "FIRMAS", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(1):
        doc.add_paragraph()
    signature = _add_legal_paragraph(doc, "__________________________________", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if responsible:
        signature.add_run(f"\n{responsible}")
    for participant in [item.strip() for item in participants.splitlines() if item.strip()]:
        name = participant.split(",", 1)[0].strip()
        if not name or name == responsible:
            continue
        doc.add_paragraph()
        participant_signature = _add_legal_paragraph(doc, "__________________________________", alignment=WD_ALIGN_PARAGRAPH.CENTER)
        participant_signature.add_run(f"\n{name}")

    target = GENERATED_DIR / safe_filename(f"ACTA-DE-REUNION-{utc_now().replace(':', '-')}.docx")
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    return target


def create_legal_docx(kind: str, fields: dict, body: str) -> Path:
    normalized_kind = kind.replace("-", "_")
    if normalized_kind == "acta_reunion":
        return _create_minutes_docx(fields, body)
    title = LEGAL_LABELS.get(normalized_kind, "ESCRITO")
    doc = Document()
    _configure_legal_document(doc)

    expediente = fields.get("expediente", "")
    sumilla = fields.get("sumilla", title.title())
    authority = fields.get("authority", "SEÑOR/A [AUTORIDAD O ENTIDAD]")
    applicant = fields.get("applicant") or fields.get("name") or "[NOMBRE COMPLETO PENDIENTE]"
    dni_value = re.sub(r"\D", "", str(fields.get("dni") or ""))
    dni = dni_value if len(dni_value) == 8 else "[DNI PENDIENTE]"
    address = fields.get("address", "[DOMICILIO]")
    email = fields.get("email", "")

    if expediente:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"EXPEDIENTE: {expediente}")
        _set_run_font(run, name="Arial", size=Pt(11), bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"SUMILLA: {sumilla}")
    _set_run_font(run, name="Arial", size=Pt(11), bold=True)

    p = doc.add_paragraph()
    run = p.add_run(authority.upper())
    _set_run_font(run, name="Arial", size=Pt(11), bold=True)

    identity = f"{applicant}, con DNI N.° {dni}, con domicilio en {address}"
    if email:
        identity += f", correo electrónico {email}"
    identity += ", ante usted me presento y digo:"
    p = doc.add_paragraph(identity)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)

    sections = _split_legal_sections(body)
    for heading, content in sections:
        p = doc.add_paragraph()
        run = p.add_run(heading.upper())
        _set_run_font(run, name="Arial", size=Pt(11), bold=True)
        for paragraph in content:
            p = doc.add_paragraph(paragraph)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.line_spacing = 1.5

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("POR TANTO:")
    _set_run_font(run, name="Arial", size=Pt(11), bold=True)
    p = doc.add_paragraph("A usted solicito admitir el presente escrito y resolver conforme a los hechos, documentos y fundamentos aplicables.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)

    city_date = fields.get("city_date", "Abancay, [FECHA]")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(city_date)

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("__________________________________\n")
    p.add_run(applicant)

    target = GENERATED_DIR / safe_filename(f"{title}-{utc_now().replace(':', '-')}.docx")
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    return target


def _split_legal_sections(body: str) -> list[tuple[str, list[str]]]:
    text = (body or "").strip()
    if not text:
        return []
    lines = text.splitlines()
    sections: list[tuple[str, list[str]]] = []
    current_heading = "I. Petitorio"
    current: list[str] = []
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        heading = re.match(r"^#{1,3}\s+(.+)$", line)
        roman_heading = re.match(r"^[IVXLCDM]+[.)-]\s+(.+)$", line, re.I)
        numbered_heading = re.match(r"^\d+[.)-]\s+(.+)$", line)
        legal_heading = roman_heading or (
            numbered_heading
            if numbered_heading and numbered_heading.group(1).isupper() and len(numbered_heading.group(1)) < 80
            else None
        )
        if heading or legal_heading or line.endswith(":") and len(line) < 80:
            if current:
                sections.append((current_heading, current))
            current_heading = line.rstrip(":")
            current = []
        else:
            current.append(line)
    if current:
        sections.append((current_heading, current))
    return sections or [("I. Contenido", [text])]
